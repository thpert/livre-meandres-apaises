from pathlib import Path
import re
import markdown
from bs4 import BeautifulSoup

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


MD_DIR = Path("md")
OUTPUT_DOCX = "fusion.docx"


def add_toc(paragraph):
    """
    Ajoute un champ Table des matières Word.
    Il faudra ensuite mettre à jour la table dans Word (F9).
    """
    run = paragraph.add_run()

    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")

    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = r'TOC \o "1-3" \h \z \u'

    fld_char_separate = OxmlElement("w:fldChar")
    fld_char_separate.set(qn("w:fldCharType"), "separate")

    text = OxmlElement("w:t")
    text.text = "Table des matières (mettre à jour avec F9)"

    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")

    run._r.append(fld_char_begin)
    run._r.append(instr_text)
    run._r.append(fld_char_separate)
    run._r.append(text)
    run._r.append(fld_char_end)


def markdown_to_docx(doc, md_content):
    html = markdown.markdown(md_content)
    soup = BeautifulSoup(html, "html.parser")

    for element in soup.children:

        if getattr(element, "name", None) is None:
            continue

        if element.name == "h1":
            doc.add_heading(element.get_text().strip(), level=1)

        elif element.name == "h2":
            doc.add_heading(element.get_text().strip(), level=2)

        elif element.name == "h3":
            doc.add_heading(element.get_text().strip(), level=3)

        elif element.name == "p":
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p.add_run(element.get_text())

        elif element.name in ("ul", "ol"):
            for li in element.find_all("li", recursive=False):
                p = doc.add_paragraph(style="List Bullet")
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                p.add_run(li.get_text())


def extract_title(lines):
    """
    Cherche le premier titre markdown.
    Sinon utilise la première ligne non vide.
    """
    for line in lines:
        s = line.strip()
        if not s:
            continue

        if s.startswith("#"):
            return s.lstrip("#").strip()

        return s

    return "Sans titre"


def numeric_sort_key(path):
    m = re.match(r"^(\d+)", path.stem)
    return int(m.group(1)) if m else float("inf")


def main():
    doc = Document()

    # Titre du document
    doc.add_heading("Document fusionné", level=0)

    # Table des matières
    toc_para = doc.add_paragraph()
    add_toc(toc_para)

    doc.add_page_break()

    md_files = sorted(
        MD_DIR.glob("*.md"),
        key=numeric_sort_key
    )

    first_file = True

    for md_file in md_files:

        if not first_file:
            doc.add_page_break()

        first_file = False

        content = md_file.read_text(encoding="utf-8")
        lines = content.splitlines()

        title = extract_title(lines)

        # Titre principal du chapitre
        # doc.add_heading(title, level=1)

        markdown_to_docx(doc, content)

    doc.save(OUTPUT_DOCX)

    print(f"Document créé : {OUTPUT_DOCX}")


if __name__ == "__main__":
    main()